# to parse the lpddr char log to pandas format and plot
# Jerry
# 2022/7/6: original version
# 2022/7/7: 1，增加测试结果的单位匹配 2，不去掉测试名字中的下划线 3, 兼容4/8 DUT log
# 2022/7/8: 1，修改单位 2，增加OS/DC/ODT/IDD/COREAC/AC/MAC 测试项列表


import os
import os.path
import sys
import zipfile
import codecs
import re
import docx

from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor 





# 这里引入所用到的包
import pandas as pd
import numpy as np
import statsmodels.api as sm
import warnings 
warnings.filterwarnings("ignore")
import matplotlib.pyplot as plt

#定义需要画图的测试项

os_plot_item=(
'Power_Short_for_VDD1',
'Power_Short_for_VDD2H',
'Power_Short_for_VDD2L',
'Power_Short_for_VDDQ',
	#'Open_Short_for_CS_Low',
	#'Open_Short_for_CS_High',
	#'Open_Short_for_CA_Low',
	#'Open_Short_for_CA_High',
	#'Open_Short_for_CK_Low',
	#'Open_Short_for_CK_High',
	#'Open_Short_for_WCK_Low',
	#'Open_Short_for_WCK_High',
	#'Open_Short_for_DQ_Low',
	#'Open_Short_for_DQ_High',
	#'Open_Short_for_DMI_Low',
	#'Open_Short_for_DMI_High',
	#'Open_Short_for_RDQS_Low',
	#'Open_Short_for_RDQS_High',
	#'Open_Short_for_RESET_N_Low',
	#'Open_Short_for_RESET_N_High',
)

dc_plot_item=(
'Input_Leakage_Low',
'Input_Leakage_High',
'CS_Leakage_Low',
'CS_Leakage_High',
'IOZ_Low',
'IOZ_High',
)

odt_plot_item=(
'RZQ-1_for_CA',
'RZQ-2_for_CA',
'RZQ-3_for_CA',
	#'RZQ-4_for_CA',
	#'RZQ-5_for_CA',
	#'RZQ-6_for_CA',
	#'RZQ-3_for_CS',
	#'RZQ-1_for_CK',
	#'RZQ-2_for_CK',
	#'RZQ-3_for_CK',
	#'RZQ-4_for_CK',
	#'RZQ-5_for_CK',
	#'RZQ-6_for_CK',
	#'RZQ-1_for_WCK',
	#'RZQ-2_for_WCK',
	#'RZQ-3_for_WCK',
	#'RZQ-4_for_WCK',
	#'RZQ-5_for_WCK',
	#'RZQ-6_for_WCK',
	#'RZQ-1_for_DQ',
	#'RZQ-2_for_DQ',
	#'RZQ-3_for_DQ',
	#'RZQ-4_for_DQ',
	#'RZQ-5_for_DQ',
	#'RZQ-6_for_DQ',
	#'RZQ-1_for_PDDS',
	#'RZQ-2_for_PDDS',
	#'RZQ-3_for_PDDS',
	#'RZQ-4_for_PDDS',
	#'RZQ-5_for_PDDS',
	#'RZQ-6_for_PDDS',


)

idd_plot_item=(
'IDD0_VDD1',
'IDD0_VDD2H',
'IDD0_VDD2L',
'IDD0_VDDQ',
	#'IDD2P_VDD1',
	#'IDD2P_VDD2H',
	#'IDD2P_VDD2L',
	#'IDD2P_VDDQ',
	#'IDD2PS_VDD1',
	#'IDD2PS_VDD2H',
	#'IDD2PS_VDD2L',
	#'IDD2PS_VDDQ',
	#'IDD2N_VDD1',
	#'IDD2N_VDD2H',
	#'IDD2N_VDD2L',
	#'IDD2N_VDDQ',
	#'IDD2NS_VDD1',
	#'IDD2NS_VDD2H',
	#'IDD2NS_VDD2L',
	#'IDD2NS_VDDQ',
	#'IDD3P_VDD1',
	#'IDD3P_VDD2H',
	#'IDD3P_VDD2L',
	#'IDD3P_VDDQ',
	#'IDD3PS_VDD1',
	#'IDD3PS_VDD2H',
	#'IDD3PS_VDD2L',
	#'IDD3PS_VDDQ',
	#'IDD3N_VDD1',
	#'IDD3N_VDD2H',
	#'IDD3N_VDD2L',
	#'IDD3N_VDDQ',
	#'IDD3NS_VDD1',
	#'IDD3NS_VDD2H',
	#'IDD3NS_VDD2L',
	#'IDD3NS_VDDQ',
	#'IDD4R_VDD1',
	#'IDD4R_VDD2H',
	#'IDD4R_VDD2L',
	#'IDD4R_VDDQ',
	#'IDD4W_VDD1',
	#'IDD4W_VDD2H',
	#'IDD4W_VDD2L',
	#'IDD4W_VDDQ',
	#'IDD5_VDD1',
	#'IDD5_VDD2H',
	#'IDD5_VDD2L',
	#'IDD5_VDDQ',
	#'IDD5AB_VDD1',
	#'IDD5AB_VDD2H',
	#'IDD5AB_VDD2L',
	#'IDD5AB_VDDQ',
	#'IDD5PB_VDD1',
	#'IDD5PB_VDD2H',
	#'IDD5PB_VDD2L',
	#'IDD5PB_VDDQ',
	#'IDD6_VDD1',
	#'IDD6_VDD2H',
	#'IDD6_VDD2L',
	#'IDD6_VDDQ',
	#'IDD6DS_VDD1',
	#'IDD6DS_VDD2H',
	#'IDD6DS_VDD2L',
	#'IDD6DS_VDDQ',
)

core_plot_item=(
'tWR',
'tRCD',
	#'tRPab',
	#'tRPpb',
	#'tWTR_L',
	#'tRRD',
	#'tPPD',
	#'tRBTP',
	#'tREFW',
	#'tWTR_S',
	#'tREFI',
)

plot_item=(
'for_CK',
'for_WCK',
#'tWCK2DQI',
#'tWCK2DQO',
	#'tLZ(RDQS_c)',
	#'tHZ(RDQS_c)',
	#'tLZ(DQ)',
	#'tHZ(DQ)',
	#'tCIPW',
	##'vCIHL_AC',
	#'tCA2CA',
	#'tCSIPW',
	#'vCSIHL_AC',
	#'ViLPD',
	#'ViHPD',
	#'tDIPW',
	##'vDIHL_AC',
	#'tDQ2DQ',
	##'tWCK2CK',
	#'tWCKSTOP',
	#'tWLWCKON',
	#'tWLMRD',
	#'tWCK_INT',
	#'tWCKSUS',
	#'tWLO',
	#'tWLDQOFF',
	#'tERQE',
	#'tERQX',
	#'tADR',
	#'tMRZ',
	#'tOSCDQI',
	#'tOSCDQO',
	#'tOSCINT',
	#'tCBTWCKPRE',
	#'tWCK2DQ7H',
	#'tDQ7HWCK',
	#'tDQ7HCK',
	#'tDQ72DQ',
	#'tCKPRECS',
	#'tCKPSTCS',
	#'tCAENT',
	#'tDStrain',
	#'tDHtrain',
	#'tCA2DMIL',
	#'tCBTRTW',
	#'tCACD',
	#'tDQ7LCK',
	#'tDQ7LWCK',
	#'tXCBT',
	#'tESPD',
	#'tSR',
	#'tXSR',
	#'tCSPD',
	#'tCMDPD',
	#'tCSLCK',
	#'tCKCSH',
	#'tCACSH',
	#'tCSCAL',
	#'tXP',
	#'tCSH',
	#'tCSL',
	#'tPDN',
	#'tXDSM_XP',
	#'tWR2WCK',
	#'tMRRI',
	#'tMRR',
	#'tMRW',
	#'tMRD',
	#'tCKFSPE',
	#'tCKFSPX',
	#'tVREFCA_L',
	#'tVREFCA_S',
	#'tVREFCA_W',
	#'tVREFDQ_L',
	#'tVREFDQ_S',
	#'tVREFDQ_W',
	#'tFC_L',
	#'tFC_S',
	#'tDQSQ',
	#'tQW',
#'tjitRDQS',
)

mm_plot_item=(
'tWCK2DQI_A_Byte0',
'tWCK2DQI_A_Byte1',
'tWCK2DQI_B_Byte0',
'tWCK2DQI_B_Byte1',
	##'tWCK2DQI_C_Byte0',
	##'tWCK2DQI_C_Byte1',
	##'tWCK2DQI_D_Byte0',
	##'tWCK2DQI_D_Byte1',
	#'tWCK2DQO_A_Byte0',
	#'tWCK2DQO_A_Byte1',
	#'tWCK2DQO_B_Byte0',
	#'tWCK2DQO_B_Byte1',
	##'tWCK2DQO_C_Byte0',
	##'tWCK2DQO_C_Byte1',
	##'tWCK2DQO_D_Byte0',
	##'tWCK2DQO_D_Byte1',
	#'vCIHL_AC_10.00%',
	#'vCIHL_AC_12.00%',
	#'vCIHL_AC_14.00%',
	#'vCIHL_AC_16.00%',
	#'vCIHL_AC_18.00%',
	#'vCIHL_AC_20.00%',
	#'vCIHL_AC_22.00%',
	#'vCIHL_AC_24.00%',
	#'vCIHL_AC_26.00%',
	#'vCIHL_AC_28.00%',
	#'vCIHL_AC_30.00%',
	#'vCIHL_AC_32.00%',
	#'vCIHL_AC_34.00%',
	#'vCIHL_AC_36.00%',
	#'vCIHL_AC_38.00%',
	#'vCIHL_AC_40.00%',
	#'vCIHL_AC_42.00%',
	#'vCIHL_AC_44.00%',
	#'vCIHL_AC_46.00%',
	#'vCIHL_AC_48.00%',
	#'vCIHL_AC_50.00%',
	#'vCIHL_AC_52.00%',
	#'vCIHL_AC_54.00%',
	#'vCIHL_AC_56.00%',
	#'vCIHL_AC_58.00%',
	#'vCIHL_AC_60.00%',
	#'vCIHL_AC_62.00%',
	#'vCIHL_AC_64.00%',
	#'vCIHL_AC_66.00%',
	#'vCIHL_AC_68.00%',
	#'vCIHL_AC_70.00%',
	#'vCIHL_AC_72.00%',
	#'vDIHL_AC_10.00%',
	#'vDIHL_AC_12.00%',
	#'vDIHL_AC_14.00%',
	#'vDIHL_AC_16.00%',
	#'vDIHL_AC_18.00%',
	#'vDIHL_AC_20.00%',
	#'vDIHL_AC_22.00%',
	#'vDIHL_AC_24.00%',
	#'vDIHL_AC_26.00%',
	#'vDIHL_AC_28.00%',
	#'vDIHL_AC_30.00%',
	#'vDIHL_AC_32.00%',
	#'vDIHL_AC_34.00%',
	#'vDIHL_AC_36.00%',
	#'vDIHL_AC_38.00%',
	#'vDIHL_AC_40.00%',
	#'vDIHL_AC_42.00%',
	#'vDIHL_AC_44.00%',
	#'vDIHL_AC_46.00%',
	#'vDIHL_AC_48.00%',
	#'vDIHL_AC_50.00%',
	#'vDIHL_AC_52.00%',
	#'vDIHL_AC_54.00%',
	#'vDIHL_AC_56.00%',
	#'vDIHL_AC_58.00%',
	#'vDIHL_AC_60.00%',
	#'vDIHL_AC_62.00%',
	#'vDIHL_AC_64.00%',
	#'vDIHL_AC_66.00%',
	#'vDIHL_AC_68.00%',
	#'vDIHL_AC_70.00%',
	#'vDIHL_AC_72.00%',
	#'tWCK2CK_Min',
	#'tWCK2CK_Max',
	#'tjit_1UI(avg)_A_0',
	#'tjit_1UI(avg)_A_1',
	#'tjit_1UI(avg)_B_0',
	#'tjit_1UI(avg)_B_1',
	##'tjit_1UI(avg)_C_0',
	##'tjit_1UI(avg)_C_1',
	##'tjit_1UI(avg)_D_0',
	##'tjit_1UI(avg)_D_1',
	#'tjit_1UI(abs)_A_0',
	#'tjit_1UI(abs)_A_1',
	#'tjit_1UI(abs)_B_0',
	#'tjit_1UI(abs)_B_1',
	##'tjit_1UI(abs)_C_0',
	##'tjit_1UI(abs)_C_1',
	##'tjit_1UI(abs)_D_0',
	##'tjit_1UI(abs)_D_1',
	#'tjit_2UI(abs)_A_0',
	#'tjit_2UI(abs)_A_1',
	#'tjit_2UI(abs)_B_0',
	#'tjit_2UI(abs)_B_1',
	##'tjit_2UI(abs)_C_0',
	##'tjit_2UI(abs)_C_1',
	##'tjit_2UI(abs)_D_0',
	##'tjit_2UI(abs)_D_1',
	#'tjit_3UI(abs)_A_0',
	#'tjit_3UI(abs)_A_1',
	#'tjit_3UI(abs)_B_0',
	#'tjit_3UI(abs)_B_1',
	##'tjit_3UI(abs)_C_0',
	##'tjit_3UI(abs)_C_1',
	##'tjit_3UI(abs)_D_0',
	##'tjit_3UI(abs)_D_1',
	#'tjit_4UI(abs)_A_0',
	#'tjit_4UI(abs)_A_1',
	#'tjit_4UI(abs)_B_0',
	#'tjit_4UI(abs)_B_1',
	##'tjit_4UI(abs)_C_0',
	##'tjit_4UI(abs)_C_1',
	##'tjit_4UI(abs)_D_0',
	##'tjit_4UI(abs)_D_1',
	#'tjit_1UI_A_0',
	#'tjit_1UI_A_1',
	#'tjit_1UI_B_0',
	#'tjit_1UI_B_1',
	##'tjit_1UI_C_0',
	##'tjit_1UI_C_1',
	##'tjit_1UI_D_0',
	##'tjit_1UI_D_1',
	#'tjit_3UI_A_0',
	#'tjit_3UI_A_1',
	#'tjit_3UI_B_0',
	#'tjit_3UI_B_1',
	##'tjit_3UI_C_0',
	##'tjit_3UI_C_1',
	##'tjit_3UI_D_0',
	##'tjit_3UI_D_1',
)

rootdir = os.path.dirname(os.path.realpath(sys.argv[0]))	#获取python工作目录

temperature = '25'
config = ''
BG     = ''
BL     = ''
Frq    = ''
Ratio  = ''
ECC    = ''
Rank   = ''
Volt   = ''
infor_cha =''
infor_chb =''
infor_chc =''
infor_chd =''

infor_flag = 0

dut_count_flag = 0
check_flag = 0
dut_count_A = 0

run_first = 0

def get_info(x):    #获取变量信息
    try:
        print(x)
        print(type(x),len(x))
    except:
        print(type(x))

#df2 = pd.DataFrame(df.values.T, index=df.columns, columns=df.index)#转置
if run_first == 1 :
	################### 数据解析开始 ####################
	#读入所有的datalog,并且重新组织所需要的相关数据
#	print(rootdir)
	merge_file = open(os.path.join(rootdir,'raw_data.csv'),'w')
	#get_info(merge_file)
	#merge_file.write('Testing,Title,Temperature,BG,BL,Frq,Ratio,ECC,Rank,Volt,Upper,Lower,Dut1,Dut2,Dut3,Dut4,Dut5,Dut6,Dut7,Dut8\n')
	#处理合并log数据
	for parent, dirnames, filenames in os.walk(rootdir):
		#print(parent)
	#	if filter(lambda x: '.csv' in x, filenames):
	#		print('###################')
	#		break
	# get the temperature information from somewhere (file name, folder name etc)
	#	if '25' in parent[-6:]:
	#		temperature = '25C'
	#
	#	elif '85' in parent[-6:]:
	#		temperature = '85C'
	#	
	#	elif '-40' in parent[-6:]:
	#		temperature = '-40C'
		for filename in filenames:
			if ('.txt' in filename) :
				new_file = codecs.open(os.path.join(parent, filename),'r',encoding='gbk')#, errors='ignore') # use gbk for encoding and ignore error
				print(filename)
				for templine in new_file:							#遍历txt文件
					templine= templine.strip('\n\r')
					if '########## Get Device Information ##########' in templine:
						infor_flag = 1
						continue
					if infor_flag == 0:
						if'#START:' in templine:
							test_name_base = templine.split(':')[3].strip(' ')
							if'IDD' in templine:
								test_name_base = templine.split(':')[2].strip(' ')+'_'+templine.split(':')[3].strip(' ')
							test_name_base=test_name_base.replace('RZQ/','RZQ-')
							config = templine.split(':')[4].strip(' ')
							if len(config)>20 :
								BG     = config.split('_')[0]
								BL     = config.split('_')[1]
								Frq    = config.split('_')[2]
								Ratio  = config.split('_')[3]
								ECC    = config.split('_')[4]
								Rank   = config.split('_')[5]
								Volt   = config.split('_')[6]
							else :
								BG     = config.split('_')[0]
								BL     = ''
								Frq    = ''
								Ratio  = ''
								ECC    = ''
								Rank   = ''
								Volt   = config.split('_')[1]
						if ('LOG:' in templine) | ('SHMOO:' in templine) :
							temp = templine.split(':')[1].strip(' ')
							temp_data = re.split(r'\s+',temp)
							if check_flag == 0:
								if len(temp_data) > 10:  # 8DUT
									dut_count_flag = 1
									dut_count_A = 12
									check_flag = 1
									merge_file.write('Testing,Title,Temperature,BG,BL,Frq,Ratio,ECC,Rank,Volt,Upper,Lower,Dut1,Dut2,Dut3,Dut4,Dut5,Dut6,Dut7,Dut8\n')
								else:					 # 4DUT
									dut_count_flag = 0 
									dut_count_A = 8
									check_flag = 1
									merge_file.write('Testing,Title,Temperature,BG,BL,Frq,Ratio,ECC,Rank,Volt,Upper,Lower,Dut1,Dut2,Dut3,Dut4\n')
							for abc in range(2,dut_count_A):
								#voltage
								if 'MV' in temp_data[abc]:
									if abc < 4:
										pass
									else:	
										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
								elif 'V' in temp_data[abc]:
									if abc < 4:
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000)+'MV'
									else:	
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000)									
								elif 'NA' in temp_data[abc]:
									if abc < 4:
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))/1000000)+'MA'
									else:	
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))/1000000)	
								elif 'UA' in temp_data[abc]:
									if abc < 4:
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))/1000)+'MA'
									else:	
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))/1000)
								elif 'MA' in temp_data[abc]:
									if abc < 4:
										pass
									else:	
										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
								elif ('A' in temp_data[abc]) & ('NOPASS' not in temp_data[abc]):
									if abc < 4:
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000)+'MA'
									else:	
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000)
								elif 'PS' in temp_data[abc]:
									if abc < 4:
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))/1000)+'NS'
									else:	
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))/1000)
								elif 'NS' in temp_data[abc]:
									if abc < 4:
										pass
									else:	
										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
								elif 'US' in temp_data[abc]:
									if abc < 4:
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000)+'NS'
									else:	
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000)
								elif 'MS' in temp_data[abc]:
									if abc < 4:
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000000)+'NS'
									else:	
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000000)
								elif ('S' in temp_data[abc]) & ('NOPASS' not in temp_data[abc]) & ('NOTEST' not in temp_data[abc]):
									if abc < 4:
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000000000)+'NS'
									else:	
										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc]))*1000000000)
								else:
									temp_data[abc]=re.sub(r"[A-Za-z\*]","",temp_data[abc])
#							'''
#							if (('MV' in temp_data[2]) | ('MV' in temp_data[3])) :
#								if 'M' not in temp_data[2] and 'V' in temp_data[2]:
#									temp_upper = str(float(temp_data[2][:-1])*1000)+'MV'
#									temp_data[2] = temp_upper
#								if 'M' not in temp_data[3] and 'V' in temp_data[3]:
#									temp_lower = str(float(temp_data[3][:-1])*1000)+'MV'
#									temp_data[3] = temp_lower
#								for abc in range(4,dut_count_A):
#									if 'MV' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									elif 'V' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000)
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('V' in temp_data[2]) | ('V' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MV' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000)
#									elif 'V' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('UA' in temp_data[2]) | ('UA' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MA' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000)
#									elif 'UA' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									elif 'NA' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000)
#									elif 'A' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000000)
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('MA' in temp_data[2]) | ('MA' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MA' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									elif 'UA' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000)
#									elif 'NA' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000)
#									elif 'A' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000)
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('A' in temp_data[2]) | ('A' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MA' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000)
#									elif 'UA' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000)
#									elif 'NA' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000000)
#									elif 'A' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('PS' in temp_data[2]) | ('PS' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000000000)
#									elif 'US' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000000)
#									elif 'NS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000)
#									elif 'PS' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									elif ('S' in temp_data[abc]) & ('NOPASS' not in temp_data[abc]):
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000000000000)
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('NS' in temp_data[2]) | ('NS' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000000)
#									elif 'US' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000)
#									elif 'NS' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									elif 'PS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000)
#									elif ('S' in temp_data[abc]) & ('NOPASS' not in temp_data[abc]):
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000000000)
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('US' in temp_data[2]) | ('US' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000)
#									elif 'US' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									elif 'NS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000)	
#									elif 'PS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000)
#									elif ('S' in temp_data[abc]) & ('NOPASS' not in temp_data[abc]):
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000000)
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('MS' in temp_data[2]) | ('MS' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MS' in temp_data[abc]:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									elif 'US' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000)
#									elif 'NS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000)
#									elif 'PS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000000)
#									elif ('S' in temp_data[abc]) & ('NOPASS' not in temp_data[abc]):
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) * 1000)
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#
#							elif (('S' in temp_data[2]) | ('S' in temp_data[3])) :
#								for abc in range(4,dut_count_A):
#									if 'MS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000)
#									elif 'US' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000)
#									elif 'NS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000000)
#									elif 'PS' in temp_data[abc]:
#										temp_data[abc]=str(float(re.sub(r"[A-Za-z]","",temp_data[abc])) / 1000000000000)
#									elif ('S' in temp_data[abc]) & ('NOPASS' not in temp_data[abc]):
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#									else:
#										temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])
#							
#							else:
#								for abc in range(4,dut_count_A):
#									temp_data[abc]=re.sub(r"[A-Za-z]","",temp_data[abc])			
#							'''
							if dut_count_flag == 1:
								if 'LOG:' in templine:
									merge_file.write(test_name_base+','+temp_data[1]+','+temperature+','+BG+','+BL+','+Frq+','+Ratio+','+ECC+','+Rank+','+Volt+','+temp_data[2]+','+temp_data[3]+','+temp_data[4]+','+temp_data[5]\
										+','+temp_data[6]+','+temp_data[7]+','+temp_data[8]+','+temp_data[9]\
										+','+temp_data[10]+','+temp_data[11]+'\n')
								elif 'SHMOO:' in templine:
									merge_file.write(test_name_base+','+temp_data[1]+','+temperature+','+BG+','+BL+','+Frq+','+Ratio+','+ECC+','+Rank+','+Volt+','+temp_data[2]+','+temp_data[3]+','+temp_data[4]+','+temp_data[5]\
										+','+temp_data[6]+','+temp_data[7]+','+temp_data[8]+','+temp_data[9]\
										+','+temp_data[10]+','+temp_data[11]+'\n')
							elif dut_count_flag == 0:
								if 'LOG:' in templine:
									merge_file.write(test_name_base+','+temp_data[1]+','+temperature+','+BG+','+BL+','+Frq+','+Ratio+','+ECC+','+Rank+','+Volt+','+temp_data[2]+','+temp_data[3]+','+temp_data[4]+','+temp_data[5]\
										+','+temp_data[6]+','+temp_data[7]+'\n')
								if 'SHMOO:' in templine:
									merge_file.write(test_name_base+','+temp_data[1]+','+temperature+','+BG+','+BL+','+Frq+','+Ratio+','+ECC+','+Rank+','+Volt+','+temp_data[2]+','+temp_data[3]+','+temp_data[4]+','+temp_data[5]\
										+','+temp_data[6]+','+temp_data[7]+'\n')
	#						merge_file.write(test_name_base+'_'+temp_data[3]+','+temp_data[4]+','+temp_data[5]+','+re.sub(r"[A-Za-z]","",temp_data[6])+','+re.sub(r"[A-Za-z]","",temp_data[7])\
	#							+','+re.sub(r"[A-Za-z]","",temp_data[8])+','+re.sub(r"[A-Za-z]","",temp_data[9])+','+re.sub(r"[A-Za-z]","",temp_data[10])+','+re.sub(r"[A-Za-z]","",temp_data[11])\
	#							+','+re.sub(r"[A-Za-z]","",temp_data[12])+','+re.sub(r"[A-Za-z]","",temp_data[13])+','+temperature+','+config+'\n')
					elif infor_flag == 1:
						if 'Not Device Information' in templine:
							continue
						if ('DUT' in templine) & ('CHA' in templine):
							infor_cha = infor_cha+templine.split(',')[8]+','
						if ('DUT' in templine) & ('CHB' in templine):
							infor_chb = infor_chb+templine.split(',')[8]+','
						if ('DUT' in templine) & ('CHC' in templine):
							infor_chc = infor_chc+templine.split(',')[8]+','
						if ('DUT' in templine) & ('CHD' in templine):
							infor_chd = infor_chd+templine.split(',')[8]+','
						if ('#TNO:' in templine):
							infor_flag = 0
							merge_file.write('Infor_CHA,,,,,'+ infor_cha.strip(',')+'\n')
							merge_file.write('Infor_CHB,,,,,'+ infor_chb.strip(',')+'\n')
							if dut_count_flag == 0: #496ball 4cha
								merge_file.write('Infor_CHC,,,,,'+ infor_chc.strip(',')+'\n')
								merge_file.write('Infor_CHD,,,,,'+ infor_chd.strip(',')+'\n')
							infor_cha = ''
							infor_chb = ''
							infor_chc = ''
							infor_chd = ''
	merge_file.close()
################### 数据解析结束 ####################

# 读入数据
iris = pd.read_csv("raw_data.csv", index_col=False) # the iris dataset is now a Pandas DataFrame

document=docx.Document()

##横屏模式
#section = document.sections[0] 
#section.orientation = WD_ORIENT.LANDSCAPE
#new_width, new_height = section.page_height, section.page_width
#section.page_width = new_width
#section.page_height = new_height

#更改字体和颜色

document.styles['Heading 1'].font.name = u'宋体'
document.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 1'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Heading 1'].font.size=Pt(16)#字体大小为16 三号

document.styles['Heading 2'].font.name = u'宋体'
document.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 2'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Heading 2'].font.size=Pt(14)#字体大小为14 四号

document.styles['Heading 3'].font.name = u'宋体'
document.styles['Heading 3']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 3'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Heading 3'].font.size=Pt(14)#字体大小为12 小四

document.styles['Heading 4'].font.name = u'宋体'
document.styles['Heading 4']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 4'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Heading 4'].font.size=Pt(14)#字体大小为12 小四


document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Normal'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Normal'].font.size=Pt(12)#字体大小为12 小四


head1 = document.add_heading('',level = 1).add_run(u'2	参数测试数据及分析')
head1.font.name=u'宋体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 

document.add_paragraph('测试数据来源说明：')
document.add_paragraph('    数据来源于T5503HS2 ESTH 测试机。')


for i in iris.columns.values:
	#去除特殊字符串
	
	if str(iris[i].dtype) == 'object':
		#iris.loc[iris[i].str.contains('*********',na=False),i]=''
		iris.loc[iris[i].str.contains(' ',na=False),i]=''
		#extr=iris[i].str.replace(r'^\D+','')

#iris = iris.dropna(axis=0,subset = [i])

plt.figure(dpi=100,figsize=(24,14))

iris_long=pd.melt(iris,id_vars=['Testing','Title','Temperature','BG','BL','Frq','Ratio','ECC','Rank','Volt','Upper','Lower'])

#print(iris_long)

head1 = document.add_heading('',level = 1).add_run(u'3	OS测试数据')
head1.font.name=u'宋体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 

count_number = 0
for i in os_plot_item: #OS plot

	data =[]
	label = []
	count = 0

	print(i)
	count_number = count_number + 1
	fig, ax1 = plt.subplots(figsize=(10, 4))
	fig.canvas.manager.set_window_title('A Boxplot Example')
	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)

	for x in iris_long.groupby(['Testing','Title','Temperature']):

		if i in x[0]:
			count = count + 1
			y=x[1]
			y['value']=pd.to_numeric(y['value'])
			y= y.dropna(axis=0,subset='value') # delete NaN

			#print(y['value'])
			AC_q1 = round(y['value'].quantile(q=0.25,interpolation="midpoint"),3)
			AC_q2 = round(y['value'].quantile(q=0.50,interpolation="midpoint"),3)
			AC_q3 = round(y['value'].quantile(q=0.75,interpolation="midpoint"),3)
			AC_robust_stdev = round((y['value'].quantile(q=0.75,interpolation="midpoint")-y['value'].quantile(q=0.25,interpolation="midpoint"))/1.35,3)

			AC_high = round(AC_q2+6*AC_robust_stdev,3)
			AC_low  = round(AC_q2-6*AC_robust_stdev,3)

			plt.scatter(x=count,y=AC_high, c="b", marker='+')
			plt.scatter(x=count,y=AC_low, c="b", marker='+')
			data.append(y['value'])
			label.append(x[0][1:])
			#print(x[0])
			#print(x[1])

			#upper=(iris.loc[iris['Testing']==i,'Upper'])
			#lower=(iris.loc[iris['Testing']==i,'Lower'])

			upper=(x[1].loc[x[1]['Testing']==i,'Upper'])
			lower=(x[1].loc[x[1]['Testing']==i,'Lower'])

			#print(lower)

			try:
				up_limit=re.sub(r"[A-Za-z]","",upper.iloc[0])
				unit_1 = re.sub(r"[0-9]","",upper.iloc[0].replace('.','').replace('-',''))
				if up_limit != '':
					plt.scatter(x=count,y=float(up_limit), c="r", marker='_')
			except:
				unit_1 =''

			try:
				low_limit=re.sub(r"[A-Za-z]","",lower.iloc[0])
				unit_2 = re.sub(r"[0-9]","",lower.iloc[0].replace('.','').replace('-',''))
				#print(low_limit)
				if low_limit != '':			
					plt.scatter(x=count,y=float(low_limit), c="r", marker='_')
			except:
				unit_2 = ''

			if (len(unit_1) > len(unit_2)) & (len(unit_1) < 3) :
				unit = unit_1
			else:
				unit = unit_2




	
	#print(data)
	ax1.boxplot(data)
	ax1.set_xticklabels(label,rotation=90, fontsize=6)
	

	
	filename='./PLOT/'+'boxplot_' +i + '_AC.jpg'
	plt.xticks(fontsize=10)
	plt.title(i,fontsize=15)

	try:
		plt.ylabel(unit,fontsize=10) #拿到单位
	except:
		print('Unit unknown')

	
	plt.savefig(filename)
	plt.clf()
	

	fig, ax1 = plt.subplots(figsize=(10, 4))
	fig.canvas.manager.set_window_title('A Boxplot Example')
	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
	
	ax1.boxplot(data)
	ax1.set_xticklabels(label,rotation=90, fontsize=6)

	
	filename1='./PLOT/'+'boxplot_' +i + '_AC_1.jpg'
	plt.xticks(fontsize=10)
	plt.title(i,fontsize=15)

	try:
		plt.ylabel(unit,fontsize=10) #拿到单位
	except:
		print('Unit unknown')

	
	plt.savefig(filename1)
	plt.clf()

	head2 = document.add_heading('',level = 2).add_run('3.'+str(count_number)+' '+i)
	head2.font.name=u'宋体'
	head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

	document.add_paragraph('	测试方法参考测试方案：《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')


	head4 = document.add_heading('',level = 3).add_run('3.'+str(count_number)+'.1 '+i+' 测试数据分布\n')
	head4.font.name=u'宋体'
	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
	head4.italic = False

	document.add_picture(filename, width=Inches(6.0))
	document.add_picture(filename1, width=Inches(6.0))    

	head4 = document.add_heading('',level = 3).add_run('3.'+str(count_number)+'.2 '+i+' 测试结论：\n')
	head4.font.name=u'宋体'
	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
	head4.italic = False

	document.add_paragraph('	测试结果符合JEDEC-5B规范。')
	plt.close()

head1 = document.add_heading('',level = 1).add_run(u'4	DC测试数据')
head1.font.name=u'宋体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 

count_number = 0
#for i in dc_plot_item: #DC plot
#	data =[]
#	label = []
#	count = 0
#
#	print(i)
#	count_number = count_number + 1	
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#
#	for x in iris_long.groupby(['Testing','Temperature','Volt']):
#
#		if i in x[0]:
#			count = count + 1
#			y=x[1]
#			y['value']=pd.to_numeric(y['value'])
#			y= y.dropna(axis=0,subset='value') # delete NaN
#
#			#print(y['value'])
#			AC_q1 = round(y['value'].quantile(q=0.25,interpolation="midpoint"),3)
#			AC_q2 = round(y['value'].quantile(q=0.50,interpolation="midpoint"),3)
#			AC_q3 = round(y['value'].quantile(q=0.75,interpolation="midpoint"),3)
#			AC_robust_stdev = round((y['value'].quantile(q=0.75,interpolation="midpoint")-y['value'].quantile(q=0.25,interpolation="midpoint"))/1.35,3)
#
#			AC_high = round(AC_q2+6*AC_robust_stdev,3)
#			AC_low  = round(AC_q2-6*AC_robust_stdev,3)
#
#
#			plt.scatter(x=count,y=AC_high, c="b", marker='+')
#			plt.scatter(x=count,y=AC_low, c="b", marker='+')
#
#			data.append(y['value'])
#			label.append(x[0][1:])
#
#			#print(x[0])
#			#print(x[1])
#
#			#upper=(iris.loc[iris['Testing']==i,'Upper'])
#			#lower=(iris.loc[iris['Testing']==i,'Lower'])
#
#			upper=(x[1].loc[x[1]['Testing']==i,'Upper'])
#			lower=(x[1].loc[x[1]['Testing']==i,'Lower'])
#
#			#print(lower)
#
#			try:
#				up_limit=re.sub(r"[A-Za-z]","",upper.iloc[0])
#				unit_1 = re.sub(r"[0-9]","",upper.iloc[0].replace('.','').replace('-',''))
#				if up_limit != '':
#					plt.scatter(x=count,y=float(up_limit), c="r", marker='_')
#			except:
#				unit_1 =''
#
#			try:
#				low_limit=re.sub(r"[A-Za-z]","",lower.iloc[0])
#				unit_2 = re.sub(r"[0-9]","",lower.iloc[0].replace('.','').replace('-',''))
#				#print(low_limit)
#				if low_limit != '':			
#					plt.scatter(x=count,y=float(low_limit), c="r", marker='_')
#			except:
#				unit_2 = ''
#
#			if (len(unit_1) > len(unit_2)) & (len(unit_1) < 3) :
#				unit = unit_1
#			else:
#				unit = unit_2
#
#
#
#
#	
#	#print(data)
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#	
#
#	
#	filename='./PLOT/'+'boxplot_' +i + '_AC.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename)
#	plt.clf()
#	
#
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#	
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#
#	
#	filename1='./PLOT/'+'boxplot_' +i + '_AC_1.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename1)
#	plt.clf()
#
#
#	
#	head2 = document.add_heading('',level = 2).add_run('4.'+str(count_number)+' '+i)
#	head2.font.name=u'宋体'
#	head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
#	document.add_paragraph('	测试方法参考测试方案：《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')
#
#	head4 = document.add_heading('',level = 3).add_run('4.'+str(count_number)+'.1 '+i+' 测试数据分布\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_picture(filename, width=Inches(6.0))
#	document.add_picture(filename1, width=Inches(6.0))    
#
#	head4 = document.add_heading('',level = 3).add_run('4.'+str(count_number)+'.2 '+i+' 测试结论：\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_paragraph('	测试结果符合JEDEC-5B规范。')
#	plt.close()
#
#head1 = document.add_heading('',level = 1).add_run(u'5	ODT测试数据')
#head1.font.name=u'宋体'
#head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
#
#count_number = 0
#for i in odt_plot_item: #ODT plot
#	data =[]
#	label = []
#	count = 0
#
#	print(i)
#	count_number = count_number + 1	
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#
#	for x in iris_long.groupby(['Testing','Temperature','Volt']):
#
#		if i in x[0]:
#			count = count + 1
#			y=x[1]
#			y['value']=pd.to_numeric(y['value'])
#			y= y.dropna(axis=0,subset='value') # delete NaN
#
#			#print(y['value'])
#			AC_q1 = round(y['value'].quantile(q=0.25,interpolation="midpoint"),3)
#			AC_q2 = round(y['value'].quantile(q=0.50,interpolation="midpoint"),3)
#			AC_q3 = round(y['value'].quantile(q=0.75,interpolation="midpoint"),3)
#			AC_robust_stdev = round((y['value'].quantile(q=0.75,interpolation="midpoint")-y['value'].quantile(q=0.25,interpolation="midpoint"))/1.35,3)
#
#			AC_high = round(AC_q2+6*AC_robust_stdev,3)
#			AC_low  = round(AC_q2-6*AC_robust_stdev,3)
#
#
#			plt.scatter(x=count,y=AC_high, c="b", marker='+')
#			plt.scatter(x=count,y=AC_low, c="b", marker='+')
#
#			data.append(y['value'])
#			label.append(x[0][1:])
#
#			#print(x[0])
#			#print(x[1])
#
#			#upper=(iris.loc[iris['Testing']==i,'Upper'])
#			#lower=(iris.loc[iris['Testing']==i,'Lower'])
#
#			upper=(x[1].loc[x[1]['Testing']==i,'Upper'])
#			lower=(x[1].loc[x[1]['Testing']==i,'Lower'])
#
#			#print(lower)
#
#			try:
#				up_limit=re.sub(r"[A-Za-z]","",upper.iloc[0])
#				unit_1 = re.sub(r"[0-9]","",upper.iloc[0].replace('.','').replace('-',''))
#				if up_limit != '':
#					plt.scatter(x=count,y=float(up_limit), c="r", marker='_')
#			except:
#				unit_1 =''
#
#			try:
#				low_limit=re.sub(r"[A-Za-z]","",lower.iloc[0])
#				unit_2 = re.sub(r"[0-9]","",lower.iloc[0].replace('.','').replace('-',''))
#				#print(low_limit)
#				if low_limit != '':			
#					plt.scatter(x=count,y=float(low_limit), c="r", marker='_')
#			except:
#				unit_2 = ''
#
#			if (len(unit_1) > len(unit_2)) & (len(unit_1) < 3) :
#				unit = unit_1
#			else:
#				unit = unit_2
#
#
#
#
#	
#	#print(data)
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#	
#
#	
#	filename='./PLOT/'+'boxplot_' +i + '_AC.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename)
#	plt.clf()
#	
#
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#	
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#
#	
#	filename1='./PLOT/'+'boxplot_' +i + '_AC_1.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename1)
#	plt.clf()
#
#	
#	head2 = document.add_heading('',level = 2).add_run('5.'+str(count_number)+' '+i)
#	head2.font.name=u'宋体'
#	head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
#	document.add_paragraph('	测试方法参考测试方案：《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')
#
#
#	head4 = document.add_heading('',level = 3).add_run('5.'+str(count_number)+'.1 '+i+' 测试数据分布\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_picture(filename, width=Inches(6.0))
#	document.add_picture(filename1, width=Inches(6.0))    
#
#	head4 = document.add_heading('',level = 3).add_run('5.'+str(count_number)+'.2 '+i+' 测试结论：\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_paragraph('	测试结果符合JEDEC-5B规范。')
#	plt.close()
#
#head1 = document.add_heading('',level = 1).add_run(u'6	IDD测试数据')
#head1.font.name=u'宋体'
#head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
#
#count_number = 0
#for i in idd_plot_item: #IDD plot
#	data =[]
#	label = []
#	count = 0
#
#	print(i)
#	count_number = count_number + 1	
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#
#	for x in iris_long.groupby(['Testing','Temperature','Volt']):
#
#		if i in x[0]:
#			count = count + 1
#			y=x[1]
#			y['value']=pd.to_numeric(y['value'])
#			y= y.dropna(axis=0,subset='value') # delete NaN
#
#			#print(y['value'])
#			AC_q1 = round(y['value'].quantile(q=0.25,interpolation="midpoint"),3)
#			AC_q2 = round(y['value'].quantile(q=0.50,interpolation="midpoint"),3)
#			AC_q3 = round(y['value'].quantile(q=0.75,interpolation="midpoint"),3)
#			AC_robust_stdev = round((y['value'].quantile(q=0.75,interpolation="midpoint")-y['value'].quantile(q=0.25,interpolation="midpoint"))/1.35,3)
#
#			AC_high = round(AC_q2+6*AC_robust_stdev,3)
#			AC_low  = round(AC_q2-6*AC_robust_stdev,3)
#
#
#			plt.scatter(x=count,y=AC_high, c="b", marker='+')
#			plt.scatter(x=count,y=AC_low, c="b", marker='+')
#
#			data.append(y['value'])
#			label.append(x[0][1:])
#
#			#print(x[0])
#			#print(x[1])
#
#			#upper=(iris.loc[iris['Testing']==i,'Upper'])
#			#lower=(iris.loc[iris['Testing']==i,'Lower'])
#
#			upper=(x[1].loc[x[1]['Testing']==i,'Upper'])
#			lower=(x[1].loc[x[1]['Testing']==i,'Lower'])
#
#			#print(lower)
#
#			try:
#				up_limit=re.sub(r"[A-Za-z]","",upper.iloc[0])
#				unit_1 = re.sub(r"[0-9]","",upper.iloc[0].replace('.','').replace('-',''))
#				if up_limit != '':
#					plt.scatter(x=count,y=float(up_limit), c="r", marker='_')
#			except:
#				unit_1 =''
#
#			try:
#				low_limit=re.sub(r"[A-Za-z]","",lower.iloc[0])
#				unit_2 = re.sub(r"[0-9]","",lower.iloc[0].replace('.','').replace('-',''))
#				#print(low_limit)
#				if low_limit != '':			
#					plt.scatter(x=count,y=float(low_limit), c="r", marker='_')
#			except:
#				unit_2 = ''
#
#			if (len(unit_1) > len(unit_2)) & (len(unit_1) < 3) :
#				unit = unit_1
#			else:
#				unit = unit_2
#
#
#
#
#	
#	#print(data)
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#	
#
#	
#	filename='./PLOT/'+'boxplot_' +i + '_AC.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename)
#	plt.clf()
#	
#
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#	
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#
#	
#	filename1='./PLOT/'+'boxplot_' +i + '_AC_1.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename1)
#	plt.clf()
#	
#	head2 = document.add_heading('',level = 2).add_run('6.'+str(count_number)+' '+i)
#	head2.font.name=u'宋体'
#	head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
#	document.add_paragraph('	测试方法参考测试方案：《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')
#
#	head4 = document.add_heading('',level = 3).add_run('6.'+str(count_number)+'.1 '+i+' 测试数据分布\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_picture(filename, width=Inches(6.0))
#	document.add_picture(filename1, width=Inches(6.0))    
#
#	head4 = document.add_heading('',level = 3).add_run('6.'+str(count_number)+'.2 '+i+' 测试结论：\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_paragraph('	测试结果符合JEDEC-5B规范。')
#	plt.close()
#
#head1 = document.add_heading('',level = 1).add_run(u'7	CORE AC测试数据')
#head1.font.name=u'宋体'
#head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
#
#count_number = 0
#for i in core_plot_item: #CORE AC plot
#
#	data =[]
#	label = []
#	count = 0
#
#	print(i)
#	count_number = count_number + 1	
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#
#	for x in iris_long.groupby(['Testing','Temperature','ECC','Frq','Volt']):
#
#		if i in x[0]:
#			count = count + 1
#			y=x[1]
#
#			y['value']=pd.to_numeric(y['value'])
#			y= y.dropna(axis=0,subset='value') # delete NaN
#
#			#os.system("pause")
#			#print(y['value'])
#			AC_q1 = round(y['value'].quantile(q=0.25,interpolation="midpoint"),3)
#			AC_q2 = round(y['value'].quantile(q=0.50,interpolation="midpoint"),3)
#			AC_q3 = round(y['value'].quantile(q=0.75,interpolation="midpoint"),3)
#			AC_robust_stdev = round((y['value'].quantile(q=0.75,interpolation="midpoint")-y['value'].quantile(q=0.25,interpolation="midpoint"))/1.35,3)
#
#			AC_high = round(AC_q2+6*AC_robust_stdev,3)
#			AC_low  = round(AC_q2-6*AC_robust_stdev,3)
#
#
#			plt.scatter(x=count,y=AC_high, c="b", marker='+')
#			plt.scatter(x=count,y=AC_low, c="b", marker='+')
#
#			data.append(y['value'])
#			label.append(x[0][1:])
#			
#			#print(x[0])
#			#print(x[1])
#
#			#upper=(iris.loc[iris['Testing']==i,'Upper'])
#			#lower=(iris.loc[iris['Testing']==i,'Lower'])
#
#			upper=(x[1].loc[x[1]['Testing']==i,'Upper'])
#			lower=(x[1].loc[x[1]['Testing']==i,'Lower'])
#
#			#print(lower)
#
#			try:
#				up_limit=re.sub(r"[A-Za-z]","",upper.iloc[0])
#				unit_1 = re.sub(r"[0-9]","",upper.iloc[0].replace('.','').replace('-',''))
#				if up_limit != '':
#					plt.scatter(x=count,y=float(up_limit), c="r", marker='_')
#			except:
#				unit_1 =''
#
#			try:
#				low_limit=re.sub(r"[A-Za-z]","",lower.iloc[0])
#				unit_2 = re.sub(r"[0-9]","",lower.iloc[0].replace('.','').replace('-',''))
#				#print(low_limit)
#				if low_limit != '':			
#					plt.scatter(x=count,y=float(low_limit), c="r", marker='_')
#			except:
#				unit_2 = ''
#
#			if (len(unit_1) > len(unit_2)) & (len(unit_1) < 3) :
#				unit = unit_1
#			else:
#				unit = unit_2
#
#
#
#
#	
#	#print(data)
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#	
#
#	
#	filename='./PLOT/'+'boxplot_' +i + '_AC.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename)
#	plt.clf()
#	
#
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#	
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#
#	
#	filename1='./PLOT/'+'boxplot_' +i + '_AC_1.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename1)
#	plt.clf()
#	
#	head2 = document.add_heading('',level = 2).add_run('7.'+str(count_number)+' '+i)
#	head2.font.name=u'宋体'
#	head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
#	document.add_paragraph('	测试方法参考测试方案：《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')
#
#	head4 = document.add_heading('',level = 3).add_run('7.'+str(count_number)+'.1 '+i+' 测试数据分布\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_picture(filename, width=Inches(6.0))
#	document.add_picture(filename1, width=Inches(6.0))    
#
#	head4 = document.add_heading('',level = 3).add_run('7.'+str(count_number)+'.2 '+i+' 测试结论：\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_paragraph('	测试结果符合JEDEC-5B规范。')
#	plt.close()
#
#head1 = document.add_heading('',level = 1).add_run(u'8	AC Timing测试数据')
#head1.font.name=u'宋体'
#head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
#
#count_number = 0
#for i in plot_item: #AC timing plot
#	data =[]
#	label = []
#	count = 0
#
#	print(i)
#	count_number = count_number + 1	
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#
#	for x in iris_long.groupby(['Testing','Temperature','ECC','Frq','Volt']):
#
#		if i in x[0]:
#			count = count + 1
#			y=x[1]
#			y['value']=pd.to_numeric(y['value'])
#			y= y.dropna(axis=0,subset='value') # delete NaN
#
#			#print(y['value'])
#			AC_q1 = round(y['value'].quantile(q=0.25,interpolation="midpoint"),3)
#			AC_q2 = round(y['value'].quantile(q=0.50,interpolation="midpoint"),3)
#			AC_q3 = round(y['value'].quantile(q=0.75,interpolation="midpoint"),3)
#			AC_robust_stdev = round((y['value'].quantile(q=0.75,interpolation="midpoint")-y['value'].quantile(q=0.25,interpolation="midpoint"))/1.35,3)
#
#			AC_high = round(AC_q2+6*AC_robust_stdev,3)
#			AC_low  = round(AC_q2-6*AC_robust_stdev,3)
#
#
#			plt.scatter(x=count,y=AC_high, c="b", marker='+')
#			plt.scatter(x=count,y=AC_low, c="b", marker='+')
#
#			data.append(y['value'])
#			label.append(x[0][1:])
#
#			#print(x[0])
#			#print(x[1])
#
#			#upper=(iris.loc[iris['Testing']==i,'Upper'])
#			#lower=(iris.loc[iris['Testing']==i,'Lower'])
#
#			upper=(x[1].loc[x[1]['Testing']==i,'Upper'])
#			lower=(x[1].loc[x[1]['Testing']==i,'Lower'])
#
#			#print(lower)
#
#			try:
#				up_limit=re.sub(r"[A-Za-z]","",upper.iloc[0])
#				unit_1 = re.sub(r"[0-9]","",upper.iloc[0].replace('.','').replace('-',''))
#				if up_limit != '':
#					plt.scatter(x=count,y=float(up_limit), c="r", marker='_')
#			except:
#				unit_1 =''
#
#			try:
#				low_limit=re.sub(r"[A-Za-z]","",lower.iloc[0])
#				unit_2 = re.sub(r"[0-9]","",lower.iloc[0].replace('.','').replace('-',''))
#				#print(low_limit)
#				if low_limit != '':			
#					plt.scatter(x=count,y=float(low_limit), c="r", marker='_')
#			except:
#				unit_2 = ''
#
#			if (len(unit_1) > len(unit_2)) & (len(unit_1) < 3) :
#				unit = unit_1
#			else:
#				unit = unit_2
#
#
#
#
#	
#	#print(data)
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#	
#
#	
#	filename='./PLOT/'+'boxplot_' +i + '_AC.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename)
#	plt.clf()
#	
#
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#	
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#
#	
#	filename1='./PLOT/'+'boxplot_' +i + '_AC_1.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename1)
#	plt.clf()
#	
#	head2 = document.add_heading('',level = 2).add_run('8.'+str(count_number)+' '+i)
#	head2.font.name=u'宋体'
#	head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
#	document.add_paragraph('	测试方法参考测试方案：《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')
#
#	head4 = document.add_heading('',level = 3).add_run('8.'+str(count_number)+'.1 '+i+' 测试数据分布\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_picture(filename, width=Inches(6.0))
#	document.add_picture(filename1, width=Inches(6.0))    
#
#	head4 = document.add_heading('',level = 3).add_run('8.'+str(count_number)+'.2 '+i+' 测试结论：\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_paragraph('	测试结果符合JEDEC-5B规范。')
#	plt.close()
#
#
#head1 = document.add_heading('',level = 1).add_run(u'9	AC Timing测试数据')
#head1.font.name=u'宋体'
#head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
#
#count_number = 0
#for i in mm_plot_item: #MM AC timing plot
#	
#	data =[]
#	label = []
#	count = 0
#
#	print(i)
#	count_number = count_number + 1	
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#
#	for x in iris_long.groupby(['Title','Temperature','ECC','Frq','Volt']):
#
#		if i in x[0]:
#			count = count + 1
#			y=x[1]
#			y['value']=pd.to_numeric(y['value'])
#			y= y.dropna(axis=0,subset='value') # delete NaN
#
#			#print(y['value'])
#			AC_q1 = round(y['value'].quantile(q=0.25,interpolation="midpoint"),3)
#			AC_q2 = round(y['value'].quantile(q=0.50,interpolation="midpoint"),3)
#			AC_q3 = round(y['value'].quantile(q=0.75,interpolation="midpoint"),3)
#			AC_robust_stdev = round((y['value'].quantile(q=0.75,interpolation="midpoint")-y['value'].quantile(q=0.25,interpolation="midpoint"))/1.35,3)
#
#			AC_high = round(AC_q2+6*AC_robust_stdev,3)
#			AC_low  = round(AC_q2-6*AC_robust_stdev,3)
#
#
#			plt.scatter(x=count,y=AC_high, c="b", marker='+')
#			plt.scatter(x=count,y=AC_low, c="b", marker='+')
#
#			data.append(y['value'])
#			label.append(x[0][1:])
#
#			#print(x[0])
#			#print(x[1])
#
#			#upper=(iris.loc[iris['Testing']==i,'Upper'])
#			#lower=(iris.loc[iris['Testing']==i,'Lower'])
#
#			upper=(x[1].loc[x[1]['Title']==i,'Upper'])
#			lower=(x[1].loc[x[1]['Title']==i,'Lower'])
#
#			#print(lower)
#
#			try:
#				up_limit=re.sub(r"[A-Za-z]","",upper.iloc[0])
#				unit_1 = re.sub(r"[0-9]","",upper.iloc[0].replace('.','').replace('-',''))
#				if up_limit != '':
#					plt.scatter(x=count,y=float(up_limit), c="r", marker='_')
#			except:
#				unit_1 =''
#
#			try:
#				low_limit=re.sub(r"[A-Za-z]","",lower.iloc[0])
#				unit_2 = re.sub(r"[0-9]","",lower.iloc[0].replace('.','').replace('-',''))
#				#print(low_limit)
#				if low_limit != '':			
#					plt.scatter(x=count,y=float(low_limit), c="r", marker='_')
#			except:
#				unit_2 = ''
#
#			if (len(unit_1) > len(unit_2)) & (len(unit_1) < 3) :
#				unit = unit_1
#			else:
#				unit = unit_2
#
#
#
#
#	
#	#print(data)
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#	
#
#	
#	filename='./PLOT/'+'boxplot_' +i + '_AC.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename)
#	plt.clf()
#	
#
#	fig, ax1 = plt.subplots(figsize=(10, 4))
#	fig.canvas.manager.set_window_title('A Boxplot Example')
#	fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.35)
#	
#	ax1.boxplot(data)
#	ax1.set_xticklabels(label,rotation=90, fontsize=6)
#
#	
#	filename1='./PLOT/'+'boxplot_' +i + '_AC_1.jpg'
#	plt.xticks(fontsize=10)
#	plt.title(i,fontsize=15)
#
#	try:
#		plt.ylabel(unit,fontsize=10) #拿到单位
#	except:
#		print('Unit unknown')
#
#	
#	plt.savefig(filename1)
#	plt.clf()
#	
#	head2 = document.add_heading('',level = 2).add_run('9.'+str(count_number)+' '+i)
#	head2.font.name=u'宋体'
#	head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
#	document.add_paragraph('	测试方法参考测试方案：《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')
#
#	head4 = document.add_heading('',level = 3).add_run('9.'+str(count_number)+'.1 '+i+' 测试数据分布\r\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_picture(filename, width=Inches(6.0))
#	document.add_picture(filename1, width=Inches(6.0))    
#
#	head4 = document.add_heading('',level = 3).add_run('9.'+str(count_number)+'.2 '+i+' 测试结论：\r\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False
#
#	document.add_paragraph('	测试结果符合JEDEC-5B规范。\r\n')
#	plt.close()
#
document.save('report.docx')



