import xlsxwriter
import time
import os
import docx

from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor 

import numpy as np
import pandas as pd
import matplotlib
import matplotlib.pyplot as plt

Title = ['HAND','TYPE','CASE','BM','BL','Freq','ECC','RANK','Volt','State','TNO','TARGET']

# rootdir = os.path.dirname(os.path.realpath(sys.argv[0]))	#获取python工作目录

def get_info(x):    #获取变量信息
    try:
        print(x,'\n',type(x),len(x))
    except:
        print(x,'\n',type(x))
    else:
        return

def Make_rule(file,filepath):
    Rule = []
    rdata = pd.read_excel(file,sheet_name=0)
    #rdata = rdata.loc[rdata['HAND']!='#FUNC']   #屏蔽多个用法('#FUNC','#DCLOG')
    #get_info(rdata)

    for i in rdata.groupby('HAND'):
        Temp = [i[0],'','','','','','','']
        Rule.append(Temp)
        Temp = []
        for j in i[1].groupby('CASE'):                #利用CASE进行分组
            Temp.append(j[0])
        for j in range(int(len(Temp)/8)):
            Rule.append(Temp[0:8])
            del Temp[0:8]
        if len(Temp)>0:
            Rule.append(Temp)

    newbook = xlsxwriter.Workbook(filepath+'Rule.xlsx')         # 生成xlsx
    wsheet = newbook.add_worksheet('PLOT')                      # 生成sheet 
    wsheet1 = newbook.add_worksheet('WORD')                      # 生成sheet 

    style = newbook.add_format({
    'font': 'calibri',      # 字体
    'border': 1,            # 边框
    'bold': False,          # 字体加粗
    "align": "left",        # 对齐方式
    "valign": "vcenter",    # 垂直居中
    "color": "black",       # 字体颜色
    'fg_color': '#FFFFFF',  # 背景色
    })
    style1 = newbook.add_format({'fg_color': '#4F81BD','font': 'calibri','border': 1,'bold': 1,})

    wsheet.write_row(0, 0, Title, style1)
    for i in range(len(Rule)):
        if '#' in Rule[i][0]:
            wsheet.write_row(i+1, 0, Rule[i], style1)    # 打印数据
        else:
            wsheet.write_row(i+1, 0, Rule[i], style)    # 打印数据
    newbook.close()  # 保存新生成的Excel

def Make_report(filepath):
    Rule = pd.read_excel(filepath+'Rule.xlsx',sheet_name='WORD')
    Rule = np.array(Rule).tolist()

    for i in range(len(Rule)):              #清除掉空值
        for j in range(len(Rule[i])):
            if type(Rule[i][j]) is float:
                del Rule[i][j:]
                break

    document=docx.Document()    #创建文档
    document.styles['Heading 1'].font.name = u'宋体'
    document.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Heading 1'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
    document.styles['Heading 1'].font.size=Pt(16)#字体大小为16 三号

    document.styles['Heading 2'].font.name = u'宋体'
    document.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Heading 2'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
    document.styles['Heading 2'].font.size=Pt(14)#字体大小为14 四号

    document.styles['Heading 3'].font.name = u'宋体'
    document.styles['Heading 3']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Heading 3'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
    document.styles['Heading 3'].font.size=Pt(14)#字体大小为12 小四

    document.styles['Heading 4'].font.name = u'宋体'
    document.styles['Heading 4']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Heading 4'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
    document.styles['Heading 4'].font.size=Pt(14)#字体大小为12 小四

    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
    document.styles['Normal'].font.size=Pt(12)#字体大小为12 小四

    biaoti1=2; biaoti2=0; biaoti3=0
    for case in Rule:
        if len(case)==0:
            break
        if '#' in case[0]:          #获取大标题
            biaoti1+=1; biaoti2=0
            head1 = document.add_heading('',level = 1).add_run(str(biaoti1)+'. '+case[0][1:])
            head1.font.name=u'宋体'
            head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
            document.add_paragraph('测试数据来源说明:')
            document.add_paragraph('    数据来源于T5503HS2 ESTH 测试机。')
            doc_flag = 0
            if case[1] == 'Group':
                doc_flag =1             #分组模式画图
        elif doc_flag==0:                       #获取内容开始画图
            for i in case:
                biaoti2+=1 ; biaoti3=0
                head2 = document.add_heading('',level = 2).add_run(str(biaoti1)+'.'+str(biaoti2)+' '+i)
                head2.font.name=u'宋体'
                head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                document.add_paragraph('    测试方法参考测试方案:《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')

                biaoti3+=1
                head4 = document.add_heading('',level = 3).add_run(str(biaoti1)+'.'+str(biaoti2)+'.'+str(biaoti3)+' 测试数据分布')
                head4.font.name=u'宋体'
                head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                #head4.italic = False

                i = i.replace("/", "_")                     #避免误识别，除号做下划线处理
                filename = filepath[:-6]+'PLOT\\boxplot_'+i+'.jpg'
                #filename1= filepath[:-6]+'\\PLOT\\boxplot1_'+i+'.jpg'
                document.add_picture(filename, width=Inches(6.0))
                #document.add_picture(filename1, width=Inches(6.0))

                biaoti3+=1
                head4 = document.add_heading('',level = 3).add_run(str(biaoti1)+'.'+str(biaoti2)+'.'+str(biaoti3)+' 测试结论：')
                head4.font.name=u'宋体'
                head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                head4.italic = False
                document.add_paragraph('	测试结果符合JEDEC-5B规范。')
                document.add_page_break()
        elif doc_flag==1:                       #获取内容开始画图
            biaoti2+=1 ; biaoti3=0
            head2 = document.add_heading('',level = 2).add_run(str(biaoti1)+'.'+str(biaoti2)+' '+case[0])
            head2.font.name=u'宋体'
            head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            document.add_paragraph('    测试方法参考测试方案:《LPDDR5样测特性分析测试方案_V4.0.pptx》以及《JESD209-5B》JEDEC LPDDR5标准。')
            biaoti3+=1
            head4 = document.add_heading('',level = 3).add_run(str(biaoti1)+'.'+str(biaoti2)+'.'+str(biaoti3)+' 测试数据分布')
            head4.font.name=u'宋体'
            head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            #head4.italic = False
            biaoti3+=1
            for i in case[1:]:
                i = i.replace("/", "_")                     #避免误识别，除号做下划线处理
                filename = filepath[:-6]+'\\PLOT\\boxplot_'+i+'.jpg'
                document.add_picture(filename, width=Inches(6.0))
            #head4 = document.add_heading('',level = 3).add_run(str(biaoti1)+'.'+str(biaoti2)+'.'+str(biaoti3)+' 局域分布')
            #head4.font.name=u'宋体'
            #head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            ##head4.italic = False
            #biaoti3+=1
            #for i in case[1:]:
            #    i = i.replace("/", "_")                     #避免误识别，除号做下划线处理
            #    filename1= filepath[:-6]+'\\PLOT\\boxplot1_'+i+'.jpg'
            #    document.add_picture(filename1, width=Inches(6.0))
            head4 = document.add_heading('',level = 3).add_run(str(biaoti1)+'.'+str(biaoti2)+'.'+str(biaoti3)+' 测试结论：')
            head4.font.name=u'宋体'
            head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            head4.italic = False
            document.add_paragraph('	测试结果符合JEDEC-5B规范。')
            document.add_page_break()
                    
    filename='D:\\DATA\\'+'Report.docx'
    document.save(filename)

def Make_report1(filepath):
    Rule = pd.read_excel(filepath+'Rule.xlsx',sheet_name='EXCEL')
    Rule = np.array(Rule).tolist()

    for i in range(len(Rule)):              #清除掉空值
        for j in range(len(Rule[i])):
            if type(Rule[i][j]) is float:
                del Rule[i][j:]
                break

    newbook = xlsxwriter.Workbook(filepath[:-6]+'Report.excel')         # 生成xlsx
    wsheet = newbook.add_worksheet('PLOT')                      # 生成sheet 

    biaoti1=2; biaoti2=0; biaoti3=0
    for case in Rule:
        if '#' in case[0]:          #获取大标题
            biaoti1+=1; biaoti2=0
            wsheet.write_row(i+1, 0, Rule[i])    # 打印数据

        elif case[0] == 'dis':      #去掉注释行
            #print(case)
            continue
        else:                       #获取内容开始画图
            for i in case:
                filename = filepath[:-6]+'\\PLOT\\boxplot_'+i+'.jpg'
                biaoti2+=1 ; biaoti3=0
                wsheet.write('A2', '向单元格插入一张图片：')
                wsheet.insert_image('B2', filename)

def Ploat(file,filepath):
    rdata = pd.read_excel(file,sheet_name=0)
    rdata = pd.melt(rdata,id_vars=list(rdata)[0:-8])
    Rule = pd.read_excel(filepath+'Rule.xlsx',sheet_name='PLOT')
    Rule = np.array(Rule).tolist()

    plotpath = filepath[0:-6]+'PLOT\\'      #创建图片存放路径
    try:
        os.makedirs(plotpath)
        print('创建图片存放路径:',plotpath)
    except:
        print('图片存放路径已存在:',plotpath)

    for i in range(len(Rule)):              #清除掉空值
        for j in range(len(Rule[i])):
            if type(Rule[i][j]) is float:
                del Rule[i][j:]
                break

    matplotlib.use('Agg')
    #aa = matplotlib.get_backend()
    #get_info(aa)
    plt.figure(dpi=100,figsize=(24,14)) #制定画板大小
    for case in Rule:               #解析规则文件
        if len(case)==0:
            continue
        if '#' in case[0]:          #获取title 分类方式
            part = case[1:]
        elif 'dis' in case or '' in case :      #去掉注释行    
            #print(case)
            continue
        else:                       #获取内容开始画图
            for i in case:
                temp  = rdata.loc[rdata['CASE']==i]     #筛选当前处理case数据，如筛选多个用('#1','#2')
                pdata = temp.groupby(part)              #分类

                count = 0
                boxdata =[]
                label = []
                print(i,part)
                #print(temp)
                fig, ax1 = plt.subplots(figsize=(10, 4))

                for j in pdata:         #按照分类，添加数据到画布
                    count = count + 1
                    y = j[1]
                    y['value']=pd.to_numeric(y['value'])
                    y= y.dropna(axis=0,subset='value')  # delete NaN

                    AC_q1 = round(y['value'].quantile(q=0.25,interpolation="midpoint"),3)      #获取分位数
                    AC_q2 = round(y['value'].quantile(q=0.50,interpolation="midpoint"),3)
                    AC_q3 = round(y['value'].quantile(q=0.75,interpolation="midpoint"),3)
                    AC_robust_stdev = round((y['value'].quantile(q=0.75,interpolation="midpoint")-y['value'].quantile(q=0.25,interpolation="midpoint"))/1.35,3)
                    AC_high = round(AC_q2+6*AC_robust_stdev,3)
                    AC_low  = round(AC_q2-6*AC_robust_stdev,3)
                    plt.scatter(x=count,y=AC_high, c="b", marker='+')          #画出分布上下限
                    plt.scatter(x=count,y=AC_low, c="b", marker='+')

                    boxdata.append(y['value'])
                    label.append(j[0])

                    upper=(j[1].loc[j[1]['CASE']==i,'UPER'])
                    lower=(j[1].loc[j[1]['CASE']==i,'LOWER'])
                    unit =(j[1].loc[j[1]['CASE']==i,'UNIT'])
                    #get_info(unit)

                    try:                                                               #画上下限
                        if upper.iloc[0] != 0:
                            plt.scatter(x=count,y=upper.iloc[0], c="r", marker='_')
                        if lower.iloc[0] != 0:
                            plt.scatter(x=count,y=lower.iloc[0], c="r", marker='_')
                    except:
                        pass

                ax1.boxplot(boxdata, sym='o', patch_artist=1)   #作图
                bott = 0
                if len(label)>4:
                    ax1.set_xticklabels(label,rotation=90)      #添加标签以及设置显示角度
                    for l in label:
                        if len(l)>bott:
                            bott = len(l)
                            fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.18*bott)
                else:
                    ax1.set_xticklabels(label)                  #添加标签以及设置显示角度

                ax1.yaxis.grid(True)                            #添加水平线
                plt.xticks(fontsize=10)                         #标签字体大小
                plt.title(i,fontsize=16)                        #图片创建标题
                try:
                    plt.ylabel(unit.iloc[0],fontsize=10)    #拿到数值的单位
                except:
                    print('not found Unit',i)               #没有单位的情况下show出数据title

                i = i.replace("/", "_")                     #避免误识别，除号做下划线处理
                filename = plotpath+'boxplot_' + i + '.jpg' #生成文件名
                plt.savefig(filename)
                plt.clf()
                plt.close('all')
                #***********  第二张局域图  ***********
                '''fig, ax1 = plt.subplots(figsize=(10, 4))        #画布
                ax1.boxplot(boxdata, sym='o', patch_artist=1)   #作图
                if len(label)>4:            #自适应显示X轴标签旋转角度
                    ax1.set_xticklabels(label,rotation=90)      #添加标签以及设置显示角度
                    bott=0  
                    for l in label:         #自适应显示X轴标签预留空间
                        if len(l)>bott:
                            bott = len(l)
                            fig.subplots_adjust(left=0.10, right=0.97, top=0.92, bottom=0.18*bott)
                else:
                    ax1.set_xticklabels(label)                  #添加标签以及设置显示角度
                ax1.yaxis.grid(True)                            #添加水平线
                plt.xticks(fontsize=10)                         #标签字体大小
                plt.title(i,fontsize=16)                        #图片创建标题
                try:
                    plt.ylabel(unit.iloc[0],fontsize=10)    #拿到数值的单位
                except:
                    print('not found Unit',i)               #没有单位的情况下show出数据title
                filename = plotpath+'boxplot1_' + i + '.jpg' #生成文件名
                plt.savefig(filename)
                plt.clf()
                plt.close('all')'''

if __name__ == '__main__':
    start = time.perf_counter()
    filepath = 'D:\\DATA\\INPUT\\'
    #plotpath = 'D:\\DATA\\PLOT\\'
    name = os.listdir(filepath)

    for i in range(len(name)):
        #print(name[i])
        DATABANK = []
        newfile = filepath[0:-6]+name[i][0:-4]+'.xlsx'
        RAW_file = filepath + name[i]

        if 'Rule.xlsx' not in name:
            Make_rule(RAW_file,filepath)
            print('\n已生成规则文件,请根据需求前往调整:路径',filepath,'\n')
            break

        if '.xlsx' not in RAW_file:
            #os.remove(RAW_file)
            continue
        elif 'Rule' in RAW_file:
            continue

        #GetData(RAW_file, DATABANK)
        #Data_unify(DATABANK)
        #PutDATA(newfile, DATABANK)
        func = input('文件名:'+RAW_file+'\n请选择功能:\n1.生成图片\n2.生成报告\n3.生成sheet报告\n')
        if eval(func) == 1:
            Ploat(RAW_file,filepath)       #已完成画图
        elif eval(func) == 2:
            Make_report(filepath)
        elif eval(func) == 3:
            Make_report1(filepath)
        else:
            print('输入错误')
        #os._exit(0)
    end = time.perf_counter()
    print('程序执行时间: %.5s秒' % (end-start))
