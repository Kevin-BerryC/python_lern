# to go through all the ADV datalog in the same folder，parse the data and plot
# Jerry
# 2022/6/7


import os
import os.path
import sys
import zipfile
import codecs
import re
import docx

from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor 

from itertools import cycle
#from bubbly.bubbly import bubbleplot 
#from __future__ import division
from plotly.offline import init_notebook_mode, iplot

# 这里引入所用到的包
import pandas as pd
import numpy as np
import statsmodels.api as sm
import warnings 
warnings.filterwarnings("ignore")
import seaborn as sns
import matplotlib.pyplot as plt
from wafer_map import wm_app

sns.set(style="white", color_codes=True)


#定义需要计算的CP测试项

cp_item=(
'FR',
'tCLL',
'tCLH',
'tDVCH',
'tCHDX',
'tCLQV',
'tCLQX_DIO0',
'tCLQX_DIO1',
'tCLQX_DIO2',
'tCLQX_DIO3',
'tSHQZ_DIO0(0->HiZ)',
'tSHQZ_DIO1(0->HiZ)',
'tSHQZ_DIO2(0->HiZ)',
'tSHQZ_DIO3(0->HiZ)',
'tSHQZ_DIO0(1->HiZ)',
'tSHQZ_DIO1(1->HiZ)',
'tSHQZ_DIO2(1->HiZ)',
'tSHQZ_DIO3(1->HiZ)',
'tSLCH',
'tCHSH',
'tSHCH',
'tCHSL',
'tSHSL',
'tHLCH',
'tCHHH',
'tHHCH',
'tCHHL',
'tHHQX_DIO0(HiZ->0)',
'tHHQX_DIO1(HiZ->0)',
'tHHQX_DIO0(HiZ->1)',
'tHHQX_DIO1(HiZ->1)',
'tHLQZ_DIO0(0->HiZ)',
'tHLQZ_DIO1(0->HiZ)',
'tHLQZ_DIO0(1->HiZ)',
'tHLQZ_DIO1(1->HiZ)',
#'VIL_CS',
#'VIL_IO',
#'VIH_CS',
#'VIH_IO',
)



count_number = 0
file_count = 0
file_count_folder = 0
lot_backup=''
lot_count = 0
test_item = ''
test_flag = 0
rootdir = os.getcwd()
temperature = '25C'
corner = 'POR'
file_header ='dut,voltage,lot,temperature,uid,'
DUT1_V1={}
DUT1_V2={}
DUT1_V3={}
DUT1_V4={}
DUT1_V5={}
DUT2_V1={}
DUT2_V2={}
DUT2_V3={}
DUT2_V4={}
DUT2_V5={}
DUT3_V1={}
DUT3_V2={}
DUT3_V3={}
DUT3_V4={}
DUT3_V5={}
DUT4_V1={}
DUT4_V2={}
DUT4_V3={}
DUT4_V4={}
DUT4_V5={}


DUT1_V1_P=''
DUT1_V2_P=''
DUT1_V3_P=''
DUT1_V4_P=''
DUT1_V5_P=''
DUT2_V1_P=''
DUT2_V2_P=''
DUT2_V3_P=''
DUT2_V4_P=''
DUT2_V5_P=''
DUT3_V1_P=''
DUT3_V2_P=''
DUT3_V3_P=''
DUT3_V4_P=''
DUT3_V5_P=''
DUT4_V1_P=''
DUT4_V2_P=''
DUT4_V3_P=''
DUT4_V4_P=''
DUT4_V5_P=''


DUT1_UID='NA'
DUT2_UID='NA'
DUT3_UID='NA'
DUT4_UID='NA'

header_write = 0
#TERS=''
#TRD=''
#TRD_ECC=''
#TPROG=''
#TPROG_ECC=''


#df2 = pd.DataFrame(df.values.T, index=df.columns, columns=df.index)#转置
#读入所有的datalog,并且重新组织所需要的相关数据


dict_lot={

'3797391':'POR',
'3797435':'PFNS',
'3797475':'PSNF',
'3797389':'POR',
'3797433':'PFNS',
'3797466':'PSNF',
'3797378':'POR',
'3797437':'PFNS',
'3797477':'PSNF',
'3797380':'POR',
'3797438':'PFNS',
'3797479':'PSNF',
}



limit_data = {}


merge_file1 = open(os.path.join(rootdir,'raw_data_AC.csv'),'a')
merge_file2 = open(os.path.join(rootdir,'raw_data_AC_limit.csv'),'a')

merge_file3 = open(os.path.join(rootdir,'raw_data_TERS.csv'),'a')
merge_file4 = open(os.path.join(rootdir,'raw_data_TRD.csv'),'a')
merge_file5 = open(os.path.join(rootdir,'raw_data_TRD_ECC.csv'),'a')
merge_file6 = open(os.path.join(rootdir,'raw_data_TPROG.csv'),'a')
merge_file7 = open(os.path.join(rootdir,'raw_data_TPROG_ECC.csv'),'a')


merge_file3.write('lot,temperature,voltage,dut,block,tERS\n')
merge_file4.write('lot,temperature,voltage,dut,block,Page,tRD\n')
merge_file5.write('lot,temperature,voltage,dut,block,Page,tRD_ECC\n')
merge_file6.write('lot,temperature,voltage,dut,block,Page,tPROG\n')
merge_file7.write('lot,temperature,voltage,dut,block,Page,tPROG_ECC\n')



print(rootdir)

'''

#处理合并CHAR log数据

for parent, dirnames, filenames in os.walk(rootdir):

	print(parent)

	file_count_folder = 0

	for filename in filenames:

		raw_data = {}
		value = {}
		start_flag = 0
		stop_flag = 0

		if '25' in filename[-3:]:
			temperature = '25C'

		elif '85' in filename[-3:]:
			temperature = '85C'
	
		elif '-40' in filename[-3:]:
			temperature = '-40C'

		if ('fsdiag' in filename):
			file_count_folder = file_count_folder + 1

			corner = filename.split('_')[0]

			line_count = 0
			new_file = codecs.open(os.path.join(parent, filename),'r',encoding='gbk')#, errors='ignore') # use gbk for encoding and ignore error
			print(filename)

			lot = re.split(r'\\',parent)[-2]

			if lot == lot_backup:
				pass
			else:
				lot_backup = lot
				lot_count = lot_count + 1

			file_count = file_count + 1
	

			for templine in new_file:
				templine = templine.strip('\r\n')
				line_count = line_count + 1
				temp_data = templine.split(' ')

				if '********** Temperature' in templine:
					temperature=templine.split(' ')[3]

				if'Contact                        ETT' in templine:
					start_flag = 1

				if ' UID=' in templine:
					if 'DUT1' in templine:
						DUT1_UID = templine.split('=')[1]
					elif 'DUT2' in templine:
						DUT2_UID = templine.split('=')[1]
					elif 'DUT3' in templine:
						DUT3_UID = templine.split('=')[1]
					elif 'DUT4' in templine:
						DUT4_UID = templine.split('=')[1]


				if '***************tRST************' in templine:
					stop_flag = 1

				if 'tERS=' in templine:
					#TERS = TERS+corner+','+temperature+','+temp_data[3].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2].split('=')[1].strip('MS')+'\n'
					merge_file3.write(corner+','+temperature+','+temp_data[3].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2].split('=')[1].strip('MS')+'\n')
				if 'tRD=' in templine:
					#TRD =     TRD+    corner+','+temperature+','+temp_data[4].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2]+','+temp_data[3].split('=')[1].strip('US')+'\n'
					merge_file4.write(corner+','+temperature+','+temp_data[4].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2]+','+temp_data[3].split('=')[1].strip('US')+'\n')
				if 'tRD_ECC=' in templine:
					#TRD_ECC = TRD_ECC+corner+','+temperature+','+temp_data[4].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2]+','+temp_data[3].split('=')[1].strip('US')+'\n'
					merge_file5.write(corner+','+temperature+','+temp_data[4].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2]+','+temp_data[3].split('=')[1].strip('US')+'\n')
				if 'tPROG=' in templine:
					#TPROG =     TPROG+corner+','+temperature+','+temp_data[3].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2]+','+temp_data[3].split('=')[1].strip('US')+'\n'
					merge_file6.write(corner+','+temperature+','+temp_data[4].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2]+','+temp_data[3].split('=')[1].strip('US')+'\n')
				if 'tPROG_ECC=' in templine:
					#TPROG_ECC = TPROG_ECC+corner+','+temperature+','+temp_data[3].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2]+','+temp_data[3].split('=')[1].strip('US')+'\n'
					merge_file7.write(corner+','+temperature+','+temp_data[4].split('=')[1][:-1]+','+ temp_data[0]+','+temp_data[1]+','+temp_data[2]+','+temp_data[3].split('=')[1].strip('US')+'\n')



				if (start_flag == 1) & (stop_flag == 0):
					if ('DUT' in templine) & ('VCC=' in templine) :

						if file_count == 1:
							if temp_data[1].split('=')[0] in file_header:
								pass
							else:
								file_header=file_header+temp_data[1].split('=')[0]+','

						if ('DUT1' in templine) & ('1.65V' in templine):
							DUT1_V1[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')
						if ('DUT1' in templine) & ('1.70V' in templine):
							DUT1_V2[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')
						if ('DUT1' in templine) & ('1.80V' in templine):
							DUT1_V3[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')
						if ('DUT1' in templine) & ('1.95V' in templine):
							DUT1_V4[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')
						if ('DUT1' in templine) & ('2.00V' in templine):
							DUT1_V5[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')

						if ('DUT2' in templine) & ('1.65V' in templine):
							DUT2_V1[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT2' in templine) & ('1.70V' in templine):
							DUT2_V2[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')  
						if ('DUT2' in templine) & ('1.80V' in templine):
							DUT2_V3[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT2' in templine) & ('1.95V' in templine):
							DUT2_V4[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')
						if ('DUT2' in templine) & ('2.00V' in templine):
							DUT2_V5[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')

						if ('DUT3' in templine) & ('1.65V' in templine):
							#print(templine)
							DUT3_V1[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT3' in templine) & ('1.70V' in templine):
							DUT3_V2[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT3' in templine) & ('1.80V' in templine):
							DUT3_V3[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')
						if ('DUT3' in templine) & ('1.95V' in templine):
							DUT3_V4[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT3' in templine) & ('2.00V' in templine):
							DUT3_V5[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 


						if ('DUT4' in templine) & ('1.65V' in templine):
							DUT4_V1[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT4' in templine) & ('1.70V' in templine):
							DUT4_V2[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT4' in templine) & ('1.80V' in templine):
							DUT4_V3[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT4' in templine) & ('1.95V' in templine):
							DUT4_V4[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV') 
						if ('DUT4' in templine) & ('2.00V' in templine):
							DUT4_V5[temp_data[1].split('=')[0]] = temp_data[1].split('=')[1].strip('NUAMnSHzV')  
											

					elif 'spec = ' in templine:
						if temp_data[0] in limit_data.keys():
							pass
						else:
							limit_data[temp_data[0]] = temp_data[4]
					else:
						continue

				if 'TEST_END_TIME:' in templine:  #每个温度测试结束打印到文件并清空flag
					start_flag=0
					stop_flag =0

					if (file_count == 1) & (header_write == 0): # header只打印一次
						merge_file1.write(file_header+'\n')
						header_write = 1

					header_name=file_header.split(',')

					for x in range(5,len(header_name)):
						if header_name[x] in DUT1_V1.keys():
							DUT1_V1_P = DUT1_V1_P +DUT1_V1[header_name[x]] + ','
						else:
							DUT1_V1_P = DUT1_V1_P + ' ,' 	

						if header_name[x] in DUT1_V2.keys():
							DUT1_V2_P = DUT1_V2_P +DUT1_V2[header_name[x]] + ','
						else:
							DUT1_V2_P = DUT1_V2_P + ' ,' 

						if header_name[x] in DUT1_V3.keys():
							DUT1_V3_P = DUT1_V3_P +DUT1_V3[header_name[x]] + ','
						else:
							DUT1_V3_P = DUT1_V3_P + ' ,' 

						if header_name[x] in DUT1_V4.keys():
							DUT1_V4_P = DUT1_V4_P +DUT1_V4[header_name[x]] + ','
						else:
							DUT1_V4_P = DUT1_V4_P + ' ,' 

						if header_name[x] in DUT1_V5.keys():
							DUT1_V5_P = DUT1_V5_P +DUT1_V5[header_name[x]] + ','
						else:
							DUT1_V5_P = DUT1_V5_P + ' ,' 

						if header_name[x] in DUT2_V1.keys():
							DUT2_V1_P = DUT2_V1_P +DUT2_V1[header_name[x]] + ','
						else:
							DUT2_V1_P = DUT2_V1_P + ' ,' 	

						if header_name[x] in DUT2_V2.keys():
							DUT2_V2_P = DUT2_V2_P +DUT2_V2[header_name[x]] + ','
						else:
							DUT2_V2_P = DUT2_V2_P + ' ,' 

						if header_name[x] in DUT2_V3.keys():
							DUT2_V3_P = DUT2_V3_P +DUT2_V3[header_name[x]] + ','
						else:
							DUT2_V3_P = DUT2_V3_P + ' ,' 

						if header_name[x] in DUT2_V4.keys():
							DUT2_V4_P = DUT2_V4_P +DUT2_V4[header_name[x]] + ','
						else:
							DUT2_V4_P = DUT2_V4_P + ' ,' 

						if header_name[x] in DUT2_V5.keys():
							DUT2_V5_P = DUT2_V5_P +DUT2_V5[header_name[x]] + ','
						else:
							DUT2_V5_P = DUT2_V5_P + ' ,' 

						if header_name[x] in DUT3_V1.keys():
							DUT3_V1_P = DUT3_V1_P +DUT3_V1[header_name[x]] + ','
						else:
							DUT3_V1_P = DUT3_V1_P + ' ,' 	

						if header_name[x] in DUT3_V2.keys():
							DUT3_V2_P = DUT3_V2_P +DUT3_V2[header_name[x]] + ','
						else:
							DUT3_V2_P = DUT3_V2_P + ' ,' 

						if header_name[x] in DUT3_V3.keys():
							DUT3_V3_P = DUT3_V3_P +DUT3_V3[header_name[x]] + ','
						else:
							DUT3_V3_P = DUT3_V3_P + ' ,' 

						if header_name[x] in DUT3_V4.keys():
							DUT3_V4_P = DUT3_V4_P +DUT3_V4[header_name[x]] + ','
						else:
							DUT3_V4_P = DUT3_V4_P + ' ,' 

						if header_name[x] in DUT3_V5.keys():
							DUT3_V5_P = DUT3_V5_P +DUT3_V5[header_name[x]] + ','
						else:
							DUT3_V5_P = DUT3_V5_P + ' ,' 

						if header_name[x] in DUT4_V1.keys():
							DUT4_V1_P = DUT4_V1_P +DUT4_V1[header_name[x]] + ','
						else:
							DUT4_V1_P = DUT4_V1_P + ' ,' 	

						if header_name[x] in DUT4_V2.keys():
							DUT4_V2_P = DUT4_V2_P +DUT4_V2[header_name[x]] + ','
						else:
							DUT4_V2_P = DUT4_V2_P + ' ,' 

						if header_name[x] in DUT4_V3.keys():
							DUT4_V3_P = DUT4_V3_P +DUT4_V3[header_name[x]] + ','
						else:
							DUT4_V3_P = DUT4_V3_P + ' ,' 

						if header_name[x] in DUT4_V4.keys():
							DUT4_V4_P = DUT4_V4_P +DUT4_V4[header_name[x]] + ','
						else:
							DUT4_V4_P = DUT4_V4_P + ' ,' 

						if header_name[x] in DUT4_V5.keys():
							DUT4_V5_P = DUT4_V5_P +DUT4_V5[header_name[x]] + ','
						else:
							DUT4_V5_P = DUT4_V5_P + ' ,' 

					
					merge_file1.write('DUT1,1.65V,'+corner+','+temperature+','+DUT1_UID+','+DUT1_V1_P+'\n')
					merge_file1.write('DUT1,1.70V,'+corner+','+temperature+','+DUT1_UID+','+DUT1_V2_P+'\n')
					merge_file1.write('DUT1,1.80V,'+corner+','+temperature+','+DUT1_UID+','+DUT1_V3_P+'\n')
					merge_file1.write('DUT1,1.95V,'+corner+','+temperature+','+DUT1_UID+','+DUT1_V4_P+'\n')
					merge_file1.write('DUT1,2.00V,'+corner+','+temperature+','+DUT1_UID+','+DUT1_V5_P+'\n')
					
					merge_file1.write('DUT2,1.65V,'+corner+','+temperature+','+DUT2_UID+','+DUT2_V1_P+'\n')
					merge_file1.write('DUT2,1.70V,'+corner+','+temperature+','+DUT2_UID+','+DUT2_V2_P+'\n')
					merge_file1.write('DUT2,1.80V,'+corner+','+temperature+','+DUT2_UID+','+DUT2_V3_P+'\n')
					merge_file1.write('DUT2,1.95V,'+corner+','+temperature+','+DUT2_UID+','+DUT2_V4_P+'\n')
					merge_file1.write('DUT2,2.00V,'+corner+','+temperature+','+DUT2_UID+','+DUT2_V5_P+'\n')
			
					merge_file1.write('DUT3,1.65V,'+corner+','+temperature+','+DUT3_UID+','+DUT3_V1_P+'\n')
					merge_file1.write('DUT3,1.70V,'+corner+','+temperature+','+DUT3_UID+','+DUT3_V2_P+'\n')
					merge_file1.write('DUT3,1.80V,'+corner+','+temperature+','+DUT3_UID+','+DUT3_V3_P+'\n')
					merge_file1.write('DUT3,1.95V,'+corner+','+temperature+','+DUT3_UID+','+DUT3_V4_P+'\n')
					merge_file1.write('DUT3,2.00V,'+corner+','+temperature+','+DUT3_UID+','+DUT3_V5_P+'\n')
			
					merge_file1.write('DUT4,1.65V,'+corner+','+temperature+','+DUT4_UID+','+DUT4_V1_P+'\n')
					merge_file1.write('DUT4,1.70V,'+corner+','+temperature+','+DUT4_UID+','+DUT4_V2_P+'\n')
					merge_file1.write('DUT4,1.80V,'+corner+','+temperature+','+DUT4_UID+','+DUT4_V3_P+'\n')
					merge_file1.write('DUT4,1.95V,'+corner+','+temperature+','+DUT4_UID+','+DUT4_V4_P+'\n')
					merge_file1.write('DUT4,2.00V,'+corner+','+temperature+','+DUT4_UID+','+DUT4_V5_P+'\n')
			
					DUT1_V1_P=''
					DUT1_V2_P=''
					DUT1_V3_P=''
					DUT1_V4_P=''
					DUT1_V5_P=''
					DUT2_V1_P=''
					DUT2_V2_P=''
					DUT2_V3_P=''
					DUT2_V4_P=''
					DUT2_V5_P=''
					DUT3_V1_P=''
					DUT3_V2_P=''
					DUT3_V3_P=''
					DUT3_V4_P=''
					DUT3_V5_P=''
					DUT4_V1_P=''
					DUT4_V2_P=''
					DUT4_V3_P=''
					DUT4_V4_P=''
					DUT4_V5_P=''
		
					DUT1_UID='NA'
					DUT2_UID='NA'
					DUT3_UID='NA'
					DUT4_UID='NA'

					DUT1_V1={}
					DUT1_V2={}
					DUT1_V3={}
					DUT1_V4={}
					DUT1_V5={}
					DUT2_V1={}
					DUT2_V2={}
					DUT2_V3={}
					DUT2_V4={}
					DUT2_V5={}
					DUT3_V1={}
					DUT3_V2={}
					DUT3_V3={}
					DUT3_V4={}
					DUT3_V5={}
					DUT4_V1={}
					DUT4_V2={}
					DUT4_V3={}
					DUT4_V4={}
					DUT4_V5={}


merge_file1.close()

limit_key=''
limit_value=''


for key in limit_data.keys():
	limit_key=limit_key+key+','
	limit_value=limit_value+limit_data[key]+','

merge_file2.write(limit_key+'\n'+limit_value+'\n')
merge_file2.close()

merge_file3.close()
merge_file4.close()
merge_file5.close()
merge_file6.close()
merge_file7.close()



lot_count=6	

	
output = 'Item'+','+'POR Min D'+','+'POR Max D'+','+'POR Min C'+','+'POR Max C'+','+'POR Min G'+','+'POR Max G'+','+'PAT Min D'+','+'PAT Max D'+','+'PAT Min C'+','+'PAT Max C'+','+'PAT Min G'+','+'PAT Max G'+'\n'

'''
# 读入数据
iris = pd.read_csv("raw_data_AC.csv", sep=',',index_col=False) # the iris dataset is now a Pandas DataFrame


#lot_count = len(iris.groupby('lot').size())



limit_cal = pd.read_csv("raw_data_AC_limit.csv", sep=',',index_col=False)


#datasheet_limit = pd.read_csv("raw_datasheet_limit.txt", sep=',',index_col=False)

#limit_cal_noUnit = pd.read_csv("limit_noUnit_CHAR.txt", sep='\t',index_col=False)


#print(limit_cal['o/s (8)'])
#print(limit_cal_noUnit['o/s (8)'])
#print(iris['o/s (8)'])

document=docx.Document()

#横屏模式
section = document.sections[0] 
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

#更改字体和颜色

document.styles['Heading 1'].font.name = u'宋体'
document.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 1'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Heading 1'].font.size=Pt(16)#字体大小为16 三号

document.styles['Heading 2'].font.name = u'宋体'
document.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 2'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Heading 2'].font.size=Pt(14)#字体大小为14 四号

document.styles['Heading 3'].font.name = u'宋体'
document.styles['Heading 3']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 3'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Heading 3'].font.size=Pt(12)#字体大小为12 小四

document.styles['Heading 4'].font.name = u'宋体'
document.styles['Heading 4']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 4'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Heading 4'].font.size=Pt(12)#字体大小为12 小四


document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Normal'].font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色
document.styles['Normal'].font.size=Pt(12)#字体大小为12 小四


head1 = document.add_heading('',level = 1).add_run(u'3	AC参数测试数据及分析')

head1.font.name=u'宋体'
head1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 

document.add_paragraph('测试数据来源说明：')
document.add_paragraph('    数据来源于CP和FT，CP都用LS801A测试仪，FT都用LS811A测试仪。CP测试为#1#2@OSE 272site，其他wafer@中山184site；FT为4 site测试。')
document.add_paragraph('通过CP的数据、结合电路及datasheet 来计算CP1和CP2的参数指标。')


for i in cp_item:
	if i in iris.columns.values:
		print(i)
	#去除非数字字符串


	#去除非数字字符串
		if str(iris[i].dtype) == 'object':
			iris.loc[iris[i].str.contains('Pass',na=False),i]=''
			iris.loc[iris[i].str.contains('Fail',na=False),i]=''
			iris.loc[iris[i].str.contains('Skip',na=False),i]=''
			iris.loc[iris[i].str.contains(' ',na=False),i]=''
			#extr=iris[i].str.replace(r'^\D+','')
	
	
	#转成数字格式
			iris[i]=pd.to_numeric(iris[i])

#		if str(limit_cal_noUnit[i].dtype) == 'object':
#
#			limit_cal_noUnit.loc[limit_cal_noUnit[i].str.contains('---',na=False),i]='00'
	
			#extr=iris[i].str.replace(r'^\D+','')
	
	
	#转成数字格式
#			limit_cal_noUnit[i]=pd.to_numeric(limit_cal_noUnit[i])
		#	print(limit_cal_noUnit[i])
		#	print(limit_cal[i])

# 看下数据前5行
#iris.head(5)

iris = iris.replace('Skip',None)
#iris = iris.dropna(axis=0,subset = [i])



plt.figure(dpi=100,figsize=(12,7))
markers={'DUT1':'o','DUT2':'v','DUT3':'d','DUT4':'s'}
colors = {125:'red',105:'pink',85:'yellow',25:'green',-40:'blue',-55:'cyan'}

for i in cp_item:
	for x in iris.groupby(['dut','temperature']):	
		y=x[1]
		
		z = y.sort_values('voltage',ascending=True, inplace=False)

		if 'VI' in i:
			for abc in range(0,len(z[i])):
				z[i].values[abc] = z[i].values[abc]/float(z['voltage'].values[abc].replace('V',''))


		plt.plot(z['voltage'], z[i], linewidth=1, color=colors[z['temperature'].values[0]], marker=markers[z['dut'].values[0]],label=str(z['dut'].values[0])+'_'+str(z['temperature'].values[0]))
	
	plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0) #指定图例位置
	plt.subplots_adjust(right=0.7)

	filename='./'+'plot_' +i.replace(r'/','_').replace(r'(','_').replace(r')','').replace(r'>','') + '.jpg'
	plt.xticks(fontsize=10)
	if 'VI' in i:
		plt.title(i+'/Vcc',fontsize=15)
	else:	
		plt.title(i,fontsize=15)
	plt.grid()	


	if 'VIL' in i:
		plt.axhline(y=0.2, c="black", ls="--", lw=2)
	if 'VIH' in i:
		plt.axhline(y=0.8, c="black", ls="--", lw=2)
	for xxx in limit_cal.columns.values:
		if xxx in i:
			plt.axhline(y=float(limit_cal[xxx][0].strip('NUAMnSHhzV')), c="black", ls="--", lw=2)
		
			try:
				plt.ylabel("".join(filter(str.isalpha, limit_cal[xxx][0])),fontsize=15) #拿到单位
			except:
				print('Unit unknown')

	plt.savefig(filename)
	plt.clf()


	head2 = document.add_heading('',level = 4).add_run(' '+i.replace('_cs','').replace('_hold',''))
	head2.font.name=u'宋体'
	head2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


	document.add_paragraph('测试数据分布')
#	head4 = document.add_heading('',level = 5).add_run('测试数据分布\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False

	document.add_picture(filename, width=Inches(6.0))
#	document.add_picture(filename1, width=Inches(8.5))

	document.add_paragraph('测试结论：\n')
#	head4 = document.add_heading('',level = 5).add_run('测试结论：\n')
#	head4.font.name=u'宋体'
#	head4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#	head4.italic = False


document.save('report.docx')
